import sys
import subprocess
import os

try:
    import docx
except ImportError:
    print("Đang cài đặt thư viện python-docx...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "--quiet"])
    import docx

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('DANH SÁCH ĐẦU MỐI LIÊN HỆ DỰ ÁN', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('Dự án: Xây dựng Hệ thống Báo cáo Quản trị & Trải nghiệm Khách hàng (Dashboard & CX/KPI Management)\n')

doc.add_heading('1. Cấu trúc thông tin liên hệ', level=1)
p = doc.add_paragraph()
p.add_run('Để dự án được triển khai thuận lợi, các hạng mục công việc cần thông tin của Đơn vị Hỗ trợ. File này cung cấp danh sách và thông tin liên lạc của các bên liên quan.\n')

doc.add_heading('2. Danh sách các đầu mối (Mẫu)', level=1)

table = doc.add_table(rows=1, cols=8)
table.style = 'Table Grid'

hdr_cells = table.rows[0].cells
headers = ['STT', 'Khối/Ban', 'Tên Đơn vị', 'Vai trò / Hạng mục hỗ trợ', 'Đầu mối Cấp nhân viên (Tên/SĐT/Email)', 'Đầu mối Quản lý - Escalation (Tên/SĐT/Email)', 'Hệ thống/Nghiệp vụ liên quan', 'Note']
for i, h in enumerate(headers):
    hdr_cells[i].text = h

data = [
    ('1', 'Khối CSKH', 'TT Chăm sóc KD', 'Xác định KPI (CSAT, NPS, Khiếu nại) & Nguồn thu thập', 'Nguyễn Văn A\n0912...\na@...', 'Trần Thị B (PGĐ)\n0913...\nb@...', 'CRM, Call Center', 'Xin API gửi OTP/SMS khảo sát'),
    ('2', 'Khối IT', 'TT Dữ liệu (Data)', 'Mapping Data & Luồng xử lý (Real-time/Batch)', 'Lê Văn C\n0909...\nc@...', 'Phạm Văn D (TP)\n0988...\nd@...', 'Data Lake, BI Dashboard', 'Thời điểm đồng bộ data 2h sáng'),
    ('3', 'Nhân sự', 'Ban Tổ chức Nhân sự', 'Báo cáo & Phân quyền user (Tập đoàn - Tỉnh)', 'Nguyễn Văn E\n0977...\ne@...', 'Lê Thị F (Trưởng Ban)\n0966...\nf@...', 'Hệ thống HR/Org Chart', 'Lấy cấu trúc sơ đồ tổ chức'),
    ('4', 'IT', 'Đội Phát triển / Vận hành', 'Thiết kế Dashboard & Cơ chế Cảnh báo (Alerts)', 'Trần Văn G\n0933...\ng@...', 'Hoàng Văn H\n0944...\nh@...', 'SMS Gateway, Server', 'Thiết kế giao diện theo Branding')
]

for item in data:
    row_cells = table.add_row().cells
    for i in range(8):
        row_cells[i].text = item[i]

output_path = r'c:\Users\caida\gds-myvnpt-wiki\sample_team\uiux-audit-team\Danh_sach_lien_he_Du_an.docx'
doc.save(output_path)
print(f"==========================================")
print(f"Đã tạo file Word thành công tại: {output_path}")
print(f"==========================================")
